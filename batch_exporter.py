"""
batch_exporter.py
Generates a .docx Word document for a content batch,
matching the format of Katie's monthly batch files exactly.

Format per post:
  [Property Name] – [Post Type]
  Post #{n} – [Title]
  Reel Text: ...      (reels only)
  Caption: ...
  [hashtags]
  Media Files: (Drive links for carousel/static, placeholder for reels)
"""

import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand colors ──────────────────────────────────────────────────────────────
COLOR_PURPLE   = RGBColor(0x6C, 0x3E, 0xF4)   # Header accent
COLOR_DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # Main text
COLOR_GRAY     = RGBColor(0x55, 0x55, 0x55)   # Sub-text
COLOR_LABEL    = RGBColor(0x33, 0x33, 0x33)   # Bold labels
ACCOUNT_COLORS = {
    "TOTG":  RGBColor(0x2E, 0x75, 0xB6),   # Blue
    "CWR":   RGBColor(0x37, 0x86, 0x5E),   # Green
    "RRE":   RGBColor(0x7B, 0x4F, 0x2E),   # Brown
    "Katie": RGBColor(0x8E, 0x44, 0xAD),   # Purple
}


def _set_run_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_horizontal_rule(doc):
    """Add a thin gray horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _account_color(account_str: str) -> RGBColor:
    s = account_str.upper()
    if "TOTG" in s or "TIPIS" in s:
        return ACCOUNT_COLORS["TOTG"]
    if "CWR" in s or "CALM" in s:
        return ACCOUNT_COLORS["CWR"]
    if "RRE" in s or "RIVER ROAD" in s:
        return ACCOUNT_COLORS["RRE"]
    if "KATIE" in s:
        return ACCOUNT_COLORS["Katie"]
    return COLOR_DARK


def generate_batch_docx(
    posts: list,
    client_name: str,
    month: str,
    year: str,
) -> bytes:
    """
    Generate a .docx file for a complete content batch.
    Returns bytes that can be streamed as a download or saved to Drive.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Default paragraph font ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_DARK

    # ── Document title ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(f"{client_name}")
    _set_run_font(run, bold=True, size=20, color=COLOR_PURPLE)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(4)
    run2 = sub_p.add_run(f"{month} {year} Content Batch  ·  {len(posts)} Posts")
    _set_run_font(run2, bold=False, size=13, color=COLOR_GRAY)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(2)
    run3 = meta_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}  ·  Platform: Instagram / TikTok")
    _set_run_font(run3, size=10, color=COLOR_GRAY, italic=True)

    _add_horizontal_rule(doc)
    doc.add_paragraph()  # spacer

    # ── Posts ─────────────────────────────────────────────────────────────────
    for i, post in enumerate(posts, 1):
        account    = post.get("account", "")
        post_type  = post.get("post_type", "Post")
        title      = post.get("title", f"Post #{i}")
        caption    = post.get("caption", "")
        hashtags   = post.get("hashtags", "")
        reel_hook  = post.get("reel_hook", "")
        vid_dir    = post.get("video_direction", "")
        img_prompt = post.get("image_prompt", "")
        angle      = post.get("angle", "")
        scheduled  = post.get("scheduled_date", "")
        is_video   = post.get("is_video", False)
        sel_imgs   = post.get("selected_images", [])
        manual_links = post.get("manual_links", [])
        video_file = post.get("video_file", "")
        is_katie   = post.get("is_katie", False)

        acc_color = _account_color(account)

        # ── Post header: "ACCOUNT – POST TYPE" ────────────────────────────
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after  = Pt(2)
        r = h1.add_run(f"{account}  –  {post_type}")
        _set_run_font(r, bold=True, size=14, color=acc_color)

        # ── Post number + title ───────────────────────────────────────────
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(0)
        h2.paragraph_format.space_after  = Pt(6)
        title_text = f"Post #{i}  –  {title}"
        if scheduled:
            title_text += f"  –  {scheduled}"
        r2 = h2.add_run(title_text)
        _set_run_font(r2, bold=True, size=12, color=COLOR_DARK)

        # ── Concept angle (small italic) ──────────────────────────────────
        if angle:
            ang_p = doc.add_paragraph()
            ang_p.paragraph_format.space_after = Pt(6)
            r_ang = ang_p.add_run(f"Concept: {angle}")
            _set_run_font(r_ang, size=10, color=COLOR_GRAY, italic=True)

        # ── REEL: hook + video direction + placeholder ────────────────────
        if is_video:
            if reel_hook:
                rh_p = doc.add_paragraph()
                rh_p.paragraph_format.space_after = Pt(4)
                lb = rh_p.add_run("🎬  Reel Hook (on-screen text):  ")
                _set_run_font(lb, bold=True, size=11, color=COLOR_LABEL)
                txt = rh_p.add_run(reel_hook)
                _set_run_font(txt, size=11)

            if vid_dir:
                vd_p = doc.add_paragraph()
                vd_p.paragraph_format.space_after = Pt(4)
                lb2 = vd_p.add_run("📹  Video Direction:  ")
                _set_run_font(lb2, bold=True, size=11, color=COLOR_LABEL)
                txt2 = vd_p.add_run(vid_dir)
                _set_run_font(txt2, size=11)

            # Video file / placeholder
            vf_p = doc.add_paragraph()
            vf_p.paragraph_format.space_after = Pt(6)
            lb3 = vf_p.add_run("🎥  Video File:  ")
            _set_run_font(lb3, bold=True, size=11, color=COLOR_LABEL)
            vf_text = video_file if video_file else "[ VIDEO PLACEHOLDER — record 15–19 sec reel per direction above ]"
            txt3 = vf_p.add_run(vf_text)
            _set_run_font(txt3, size=11, italic=not bool(video_file), color=COLOR_GRAY if not video_file else None)

        # ── Caption ───────────────────────────────────────────────────────
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.space_after = Pt(2)
        cap_lb = cap_p.add_run("Caption:  ")
        _set_run_font(cap_lb, bold=True, size=11, color=COLOR_LABEL)

        if caption:
            cap_body = doc.add_paragraph()
            cap_body.paragraph_format.left_indent  = Inches(0.2)
            cap_body.paragraph_format.space_after  = Pt(2)
            cap_run = cap_body.add_run(caption)
            _set_run_font(cap_run, size=11)

        # ── Hashtags ──────────────────────────────────────────────────────
        if hashtags:
            ht_p = doc.add_paragraph()
            ht_p.paragraph_format.left_indent  = Inches(0.2)
            ht_p.paragraph_format.space_after  = Pt(8)
            ht_run = ht_p.add_run(hashtags)
            _set_run_font(ht_run, size=10, color=COLOR_GRAY)

        # ── Media: Drive links for carousel/static ────────────────────────
        if not is_video:
            media_p = doc.add_paragraph()
            media_p.paragraph_format.space_after = Pt(2)
            med_lb = media_p.add_run("📸  Media Files:  ")
            _set_run_font(med_lb, bold=True, size=11, color=COLOR_LABEL)

            links_to_show = []
            if sel_imgs:
                links_to_show = [
                    img.get("webViewLink") or img.get("name", f"Image {j+1}")
                    for j, img in enumerate(sel_imgs)
                ]
            elif manual_links:
                links_to_show = manual_links

            if links_to_show:
                for j, link in enumerate(links_to_show, 1):
                    lp = doc.add_paragraph()
                    lp.paragraph_format.left_indent = Inches(0.3)
                    lp.paragraph_format.space_after = Pt(1)
                    lr = lp.add_run(f"{j}.  {link}")
                    _set_run_font(lr, size=10, color=RGBColor(0x1A, 0x6F, 0xC4))
            else:
                # Image direction as placeholder
                if img_prompt:
                    ip = doc.add_paragraph()
                    ip.paragraph_format.left_indent = Inches(0.3)
                    ip.paragraph_format.space_after = Pt(4)
                    ir = ip.add_run(f"[ IMAGE DIRECTION: {img_prompt} ]")
                    _set_run_font(ir, size=10, italic=True, color=COLOR_GRAY)
                else:
                    ph = doc.add_paragraph()
                    ph.paragraph_format.left_indent = Inches(0.3)
                    ph.paragraph_format.space_after = Pt(4)
                    phr = ph.add_run("[ Add image Drive links here ]")
                    _set_run_font(phr, size=10, italic=True, color=COLOR_GRAY)

        # ── Separator ─────────────────────────────────────────────────────
        doc.add_paragraph()
        _add_horizontal_rule(doc)

    # ── Footer summary ────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot_p.add_run(f"End of {month} {year} Batch  ·  {len(posts)} Posts Total  ·  Generated by Censational Social Media Manager")
    _set_run_font(fr, size=9, italic=True, color=COLOR_GRAY)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
